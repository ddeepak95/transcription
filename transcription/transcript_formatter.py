from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document


def format_seconds(seconds: float | int | None) -> str:
    if seconds is None:
        return "unknown"
    total_ms = int(float(seconds) * 1000)
    hours = total_ms // 3_600_000
    minutes = (total_ms % 3_600_000) // 60_000
    secs = (total_ms % 60_000) // 1000
    millis = total_ms % 1000
    return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"


def build_transcript_text(full_response: dict[str, Any], fallback_transcript: str = "") -> str:
    diarized = full_response.get("diarized_transcript", {})
    entries = diarized.get("entries", []) if isinstance(diarized, dict) else []
    if entries:
        lines: list[str] = []
        for entry in entries:
            text = str(entry.get("transcript", "")).strip()
            if not text:
                continue
            speaker = entry.get("speaker_id", "unknown")
            start_ts = format_seconds(entry.get("start_time_seconds"))
            end_ts = format_seconds(entry.get("end_time_seconds"))
            lines.append(f"[{start_ts} - {end_ts}] Speaker {speaker}: {text}")
        if lines:
            return "\n".join(lines)

    transcript = (
        full_response.get("transcript")
        or fallback_transcript
        or ""
    )
    return str(transcript).strip()


def build_transcript_file_content(audio_name: str, full_response: dict[str, Any], fallback_transcript: str = "") -> str:
    body = build_transcript_text(full_response, fallback_transcript=fallback_transcript)
    return f"Audio: {audio_name}\n\n{body}".strip() + "\n"


def default_transcript_output_path(input_json_path: Path, audio_name: str | None = None) -> Path:
    base_name = audio_name if audio_name else input_json_path.stem
    return input_json_path.with_name(f"{base_name}_transcript.txt")


def default_transcript_docx_output_path(input_json_path: Path, audio_name: str | None = None) -> Path:
    return default_transcript_output_path(input_json_path, audio_name=audio_name).with_suffix(".docx")


def write_transcript_docx(
    output_path: Path,
    audio_name: str,
    full_response: dict[str, Any],
    fallback_transcript: str = "",
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_paragraph(f"Audio: {audio_name}")
    document.add_paragraph("")

    transcript_body = build_transcript_text(full_response, fallback_transcript=fallback_transcript)
    if transcript_body:
        for line in transcript_body.splitlines():
            document.add_paragraph(line)
    else:
        document.add_paragraph("")

    document.save(str(output_path))
